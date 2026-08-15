"""Generates the Phase 1 sample/test data pack into data/sample_docs/.

All content is synthetic/fictional — no real students, teachers, or
institutions. Re-run any time to regenerate: `python scripts/generate_sample_data.py`
(seeded, so output is deterministic).

Covers the four formats the Phase 1 task calls for (docs/final-task-assignment.md):
Excel results, a PDF, a DOCX, and test images for the Vision/OCR agent.
"""

import random
from pathlib import Path

from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from openpyxl import Workbook
from openpyxl.styles import Font
from PIL import Image, ImageDraw, ImageFilter, ImageFont

SEED = 42
OUT_DIR = Path(__file__).resolve().parents[1] / "data" / "sample_docs"
FONT_DIR = Path(r"C:\Windows\Fonts")

DEPARTMENTS = ["Computer Science", "Electrical Engineering", "Business Administration"]
COURSES = {
    "Computer Science": ["Data Structures", "Operating Systems", "Machine Learning"],
    "Electrical Engineering": ["Circuit Theory", "Signals & Systems"],
    "Business Administration": ["Financial Accounting", "Marketing Principles"],
}
TERMS = ["Spring 2025", "Fall 2025"]
FIRST_NAMES = [
    "Amina", "Rakib", "Nusrat", "Farhan", "Tanvir", "Mahia", "Sabbir", "Ishrat",
    "Kamrul", "Priya", "Sadia", "Arif", "Nabila", "Shakil", "Meherun", "Tariq",
]
LAST_NAMES = [
    "Hossain", "Rahman", "Islam", "Chowdhury", "Akter", "Karim", "Uddin", "Begum",
]


def generate_excel() -> Path:
    random.seed(SEED)
    wb = Workbook()
    ws = wb.active
    ws.title = "Results"

    headers = ["enrollment_no", "student_name", "department", "course", "term", "score"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)

    for i in range(1, 31):
        department = random.choice(DEPARTMENTS)
        course = random.choice(COURSES[department])
        name = f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"
        ws.append([
            f"ENR-2025-{i:04d}",
            name,
            department,
            course,
            random.choice(TERMS),
            random.randint(42, 98),
        ])

    for column_cells in ws.columns:
        width = max(len(str(cell.value)) for cell in column_cells) + 2
        ws.column_dimensions[column_cells[0].column_letter].width = width

    path = OUT_DIR / "enrollment_results.xlsx"
    wb.save(path)
    return path


def generate_pdf() -> Path:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Course Syllabus: Data Structures & Algorithms", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "", 11)
    pdf.ln(2)
    pdf.multi_cell(
        0, 6,
        "Department: Computer Science | Term: Spring 2025 | Credits: 3\n"
        "Instructor: Dr. A. Rahman\n\n"
        "Course Objectives:\n"
        "This course introduces fundamental data structures (arrays, linked lists, "
        "trees, graphs, hash tables) and algorithm design techniques (recursion, "
        "divide-and-conquer, dynamic programming), with emphasis on complexity analysis.",
    )
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Weekly Schedule", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 10)
    schedule = [
        ("Week 1-2", "Arrays, Linked Lists, Complexity Analysis"),
        ("Week 3-4", "Stacks, Queues, Recursion"),
        ("Week 5-6", "Trees, Binary Search Trees, Heaps"),
        ("Week 7-8", "Graphs, Traversals, Shortest Paths"),
        ("Week 9-10", "Hash Tables, Dynamic Programming"),
        ("Week 11-12", "Sorting Algorithms, Project Work"),
    ]
    for week, topic in schedule:
        pdf.cell(35, 7, week, border=1)
        pdf.cell(0, 7, topic, border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Grading Breakdown", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 10)
    for component, weight in [
        ("Assignments", "20%"), ("Midterm Exam", "25%"),
        ("Final Project", "25%"), ("Final Exam", "30%"),
    ]:
        pdf.cell(60, 7, component, border=1)
        pdf.cell(0, 7, weight, border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    path = OUT_DIR / "course_syllabus.pdf"
    pdf.output(str(path))
    return path


def generate_docx() -> Path:
    doc = Document()
    doc.add_heading("Computer Science Department — Term Report", level=1)
    doc.add_paragraph(
        "Spring 2025 term summary, prepared for the departmental review. "
        "All figures below are synthetic test data."
    )

    doc.add_heading("Enrollment Summary", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Program", "Enrolled Students"
    for program, count in [
        ("BSc Computer Science", "142"),
        ("BSc Electrical Engineering", "98"),
        ("BBA", "176"),
    ]:
        row = table.add_row().cells
        row[0].text, row[1].text = program, count

    doc.add_heading("Notes", level=2)
    for note in [
        "Average TPE (Teacher Performance Evaluation) score improved 4% over last term.",
        "Two course sections were merged due to low enrollment.",
        "Lab equipment upgrade completed ahead of schedule.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(note).font.size = Pt(11)

    path = OUT_DIR / "department_report.docx"
    doc.save(path)
    return path


def _render_text_image(lines: list[str], font_path: Path, font_size: int,
                        size: tuple[int, int]) -> Image.Image:
    img = Image.new("L", size, color=250)
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype(str(font_path), font_size)
    y = 30
    for line in lines:
        draw.text((30, y), line, fill=20, font=font)
        y += font_size + 14
    return img


def _simulate_scan(img: Image.Image, rotate_deg: float, noise_seed: int) -> Image.Image:
    random.seed(noise_seed)
    img = img.rotate(rotate_deg, expand=True, fillcolor=250)
    img = img.filter(ImageFilter.GaussianBlur(radius=0.6))
    pixels = img.load()
    w, h = img.size
    for _ in range(int(w * h * 0.01)):
        x, y = random.randrange(w), random.randrange(h)
        pixels[x, y] = random.randint(180, 255)
    return img


def generate_images() -> list[Path]:
    paths = []

    sheet = _render_text_image(
        lines=[
            "RESULT SHEET - Spring 2025",
            "",
            "ENR-2025-0007  Rakib Hossain      Data Structures     87",
            "ENR-2025-0012  Nusrat Chowdhury    Data Structures     91",
            "ENR-2025-0019  Farhan Islam        Data Structures     76",
            "ENR-2025-0023  Sadia Karim         Data Structures     94",
        ],
        font_path=FONT_DIR / "arial.ttf", font_size=22, size=(900, 320),
    )
    sheet = _simulate_scan(sheet, rotate_deg=1.5, noise_seed=SEED)
    p1 = OUT_DIR / "scanned_result_sheet.png"
    sheet.save(p1)
    paths.append(p1)

    note = _render_text_image(
        lines=[
            "Attendance Note - Week 6",
            "",
            "Section B moved to Room 204",
            "due to lab maintenance. Please",
            "inform students before Monday.",
        ],
        font_path=FONT_DIR / "ariali.ttf", font_size=24, size=(700, 260),
    )
    note = _simulate_scan(note, rotate_deg=-2.2, noise_seed=SEED + 1)
    p2 = OUT_DIR / "scanned_attendance_note.png"
    note.save(p2)
    paths.append(p2)

    return paths


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = [generate_excel(), generate_pdf(), generate_docx(), *generate_images()]
    print(f"Generated {len(generated)} files in {OUT_DIR}:")
    for path in generated:
        print(f"  - {path.name}")


if __name__ == "__main__":
    main()
