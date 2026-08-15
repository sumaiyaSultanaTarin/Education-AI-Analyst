"""DOCX text/table extraction, used by the Document Ingestion Agent."""

from docx import Document


def extract_docx_content(path: str) -> dict:
    """Extract headings/paragraphs and tables from a DOCX file.

    Returns {"paragraphs": list[str], "tables": list[list[list[str]]]}.
    Empty paragraphs (spacing artifacts) are dropped.
    """
    doc = Document(path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    tables = []
    for table in doc.tables:
        tables.append([[cell.text for cell in row.cells] for row in table.rows])

    return {"paragraphs": paragraphs, "tables": tables}
